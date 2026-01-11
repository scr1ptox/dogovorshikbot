<<<<<<< HEAD
# docx_generator.py
from pathlib import Path
=======
# noinspection PyProtectedMember,PyBroadException,PyTypeChecker,DuplicatedCode
# docx_generator.py
>>>>>>> 03ccfeb (Fix docx generation, disable PDF, stabilize template formatting)
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
<<<<<<< HEAD


=======
from pathlib import Path
from paths import TEMPLATE_CONTRACT, TEMPLATE_SCHEDULE
import subprocess
from docx.document import Document as DocxDocument
import copy

# ВАЖНО: чтобы документ оставался 1:1 как шаблон, мы НЕ трогаем стили/шрифты/абзацы.
# Генератор делает только замену плейсхолдеров, сохраняя исходное форматирование шаблона.
PRESERVE_TEMPLATE_FORMAT = True

def _clone_run_rpr(src_run, dst_run) -> None:
    """
    Копирует XML-свойства символов (w:rPr) из src_run в dst_run.
    Нужно, чтобы при жёсткой замене плейсхолдера новый run унаследовал форматирование шаблона.
    """
    # noinspection PyProtectedMember
    src_r = src_run._element
    # noinspection PyProtectedMember
    dst_r = dst_run._element

    src_rpr = src_r.find(qn('w:rPr'))
    if src_rpr is None:
        return

    dst_rpr = dst_r.find(qn('w:rPr'))
    if dst_rpr is not None:
        dst_r.remove(dst_rpr)

    dst_r.insert(0, copy.deepcopy(src_rpr))

# noinspection DuplicatedCode
>>>>>>> 03ccfeb (Fix docx generation, disable PDF, stabilize template formatting)
def _force_font(run, pt=10, family="Times New Roman"):
    # python-docx high-level API
    run.font.size = Pt(pt)
    run.font.name = family

    # clear features that visually shrink glyphs
    try:
        run.font.superscript = False
        run.font.subscript = False
        run.font.small_caps = False
        run.font.all_caps = False
<<<<<<< HEAD
=======
    # noinspection PyBroadException
>>>>>>> 03ccfeb (Fix docx generation, disable PDF, stabilize template formatting)
    except Exception:
        pass

    # noinspection PyProtectedMember
    # Low-level XML to avoid Word reapplying theme/eastAsia/half-point size:
    r = run._element
    # ensure rPr exists
    rpr = r.get_or_add_rPr()

    # ensure rFonts exists
    if rpr.rFonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    else:
        rfonts = rpr.rFonts
    rfonts.set(qn("w:ascii"), family)
    rfonts.set(qn("w:hAnsi"), family)
    rfonts.set(qn("w:eastAsia"), family)
    rfonts.set(qn("w:cs"), family)
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:csTheme"):
        if rfonts.get(qn(attr)) is not None:
            del rfonts.attrib[qn(attr)]

    # force size in half-points: 10pt => 20
    half_points = str(int(pt * 2))

    # w:sz
    sz = rpr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rpr.append(sz)
    sz.set(qn('w:val'), half_points)

    # w:szCs (complex scripts)
    szcs = rpr.find(qn('w:szCs'))
    if szcs is None:
        szcs = OxmlElement('w:szCs')
        rpr.append(szcs)
    szcs.set(qn('w:val'), half_points)

    # strip XML attrs that may make text look smaller or raised/lowered
    for tag in ("w:vertAlign", "w:position", "w:spacing", "w:smallCaps", "w:caps"):
        node = rpr.find(qn(tag))
        if node is not None:
            rpr.remove(node)

    # force baseline (explicitly cancel superscript/subscript)
    va = OxmlElement('w:vertAlign')
    va.set(qn('w:val'), 'baseline')
    rpr.append(va)

    # remove character style reference (w:rStyle) to avoid inherited small font
    rstyle = rpr.find(qn('w:rStyle'))
    if rstyle is not None:
        rpr.remove(rstyle)

    # Ensure remove theme overrides completely
    if rpr.find(qn('w:sz')) is not None:
        rpr.find(qn('w:sz')).set(qn('w:val'), half_points)
    if rpr.find(qn('w:szCs')) is not None:
        rpr.find(qn('w:szCs')).set(qn('w:val'), half_points)

def _replace_placeholder_in_paragraph_strict(paragraph, key: str, value: str, pt: int = 10, family: str = "Times New Roman") -> bool:
    """
    Жёсткая замена плейсхолдера key на value в абзаце, даже если key разбит на несколько runs.
    Алгоритм:
      1) Собираем полный текст абзаца и находим позицию key.
      2) Определяем границы по символам, какие runs затронуты.
      3) Перестраиваем runs: оставляем префикс в первом run, вставляем НОВЫЙ run со значением (принудительно Times NR 10pt),
         в последнем run оставляем только суффикс; промежуточные – очищаем.
    Возвращает True если была произведена замена.
    """
    full = "".join(r.text or "" for r in paragraph.runs)
    idx = full.find(key)
    if idx < 0:
        return False

    # Построим карту соответствия: для каждого run знаем его диапазон в full
    spans = []
    pos = 0
    for r in paragraph.runs:
        txt = r.text or ""
        spans.append((pos, pos + len(txt), r))
        pos += len(txt)

    start = idx
    end = idx + len(key)

    # Находим первый и последний run, которые пересекают плейсхолдер
    first_i = next(i for i, (a, b, _) in enumerate(spans) if not (b <= start or a >= end))
    last_i = next(i for i in range(len(spans) - 1, -1, -1) if not (spans[i][1] <= start or spans[i][0] >= end))

    # Префикс/суффикс
    first_run = spans[first_i][2]
    last_run = spans[last_i][2]

    # вычислим локальные индексы в первом и последнем run
    first_a, first_b, _ = spans[first_i]
    last_a, last_b, _ = spans[last_i]
    first_prefix_len = max(0, start - first_a)
    last_suffix_len = max(0, last_b - end)

    # 1) Первый run: оставляем только префикс до key
<<<<<<< HEAD
=======
    # Сохраняем донор форматирования ДО изменения текста (важно для подчеркиваний/линий в шаблоне)
    donor_run = first_run

>>>>>>> 03ccfeb (Fix docx generation, disable PDF, stabilize template formatting)
    first_text = first_run.text or ""
    first_run.text = first_text[:first_prefix_len]

    # 2) Все промежуточные run очищаем
    for j in range(first_i + 1, last_i):
        spans[j][2].text = ""

    # 3) Последний run: оставляем суффикс после key
    last_text = last_run.text or ""
    last_run.text = last_text[len(last_text) - last_suffix_len:] if last_suffix_len > 0 else ""

    # 4) Вставляем НОВЫЙ run со значением непосредственно после first_run
    insert_after = first_run
    new_run = paragraph.add_run("")  # добавим в конец...
    # noinspection PyProtectedMember
    paragraph._p.insert(paragraph._p.index(insert_after._r) + 1, new_run._r)
    new_run.text = value
<<<<<<< HEAD
    _force_font(new_run, pt, family)

    # Нормализуем все затронутые runs на случай странных тем/глифов
    for j in range(first_i, last_i + 1):
        _force_font(spans[j][2], pt, family)

    _normalize_paragraph(paragraph, pt, family)
=======

    if PRESERVE_TEMPLATE_FORMAT:
        # Наследуем форматирование из run, который содержал плейсхолдер в шаблоне.
        # Это критично для строк с подчеркиванием/линиями (часто завязано на rPr конкретного run).
        _clone_run_rpr(donor_run, new_run)
        # На всякий случай унаследуем и run.style (если задан)
        try:
            new_run.style = donor_run.style
        except Exception:
            pass
    else:
        _force_font(new_run, pt, family)
        # Нормализуем все затронутые runs на случай странных тем/глифов
        for j in range(first_i, last_i + 1):
            _force_font(spans[j][2], pt, family)
        _normalize_paragraph(paragraph, pt, family)
>>>>>>> 03ccfeb (Fix docx generation, disable PDF, stabilize template formatting)

    return True

def _replace_in_paragraph(paragraph, mapping: dict, fallback_font_pt: int = 10):
    # 1) Простые замены внутри отдельных runs
    changed = False
    for run in paragraph.runs:
        original = run.text or ""
        new_text = original
        for key, val in mapping.items():
            if key in new_text:
                new_text = new_text.replace(key, str(val))
        if new_text != original:
            run.text = new_text
<<<<<<< HEAD
            _force_font(run, fallback_font_pt, "Times New Roman")
=======
            if not PRESERVE_TEMPLATE_FORMAT:
                _force_font(run, fallback_font_pt, "Times New Roman")
>>>>>>> 03ccfeb (Fix docx generation, disable PDF, stabilize template formatting)
            changed = True

    # 2) Жёсткая замена, если плейсхолдеры разбиты на несколько runs
    for key, val in mapping.items():
        if key in (paragraph.text or ""):
            if _replace_placeholder_in_paragraph_strict(paragraph, key, str(val), fallback_font_pt, "Times New Roman"):
                changed = True

    # Если что-то меняли – дополнительно нормализуем все runs абзаца (на случай тем/глифов)
<<<<<<< HEAD
    if changed:
=======
    if changed and not PRESERVE_TEMPLATE_FORMAT:
>>>>>>> 03ccfeb (Fix docx generation, disable PDF, stabilize template formatting)
        _normalize_paragraph(paragraph, fallback_font_pt, "Times New Roman")

def _clear_paragraph_char_props(paragraph):
    """
    Очищает только проблемные символьные свойства на уровне абзаца (w:pPr/w:rPr),
    не удаляя весь rPr (иначе теряется формат нумерации).
    """
<<<<<<< HEAD
=======
    # noinspection PyProtectedMember
>>>>>>> 03ccfeb (Fix docx generation, disable PDF, stabilize template formatting)
    p = paragraph._element
    pPr = p.find(qn('w:pPr'))
    if pPr is None:
        return
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        return
    # удалить узлы, которые визуально "мельчат" или сдвигают текст
    for tag in ("w:vertAlign", "w:position", "w:spacing", "w:smallCaps", "w:caps", "w:rStyle"):
        node = rPr.find(qn(tag))
        if node is not None:
            rPr.remove(node)
    # Явно вернуть baseline
    va = OxmlElement('w:vertAlign')
    va.set(qn('w:val'), 'baseline')
    rPr.append(va)

def _ensure_para_numbering_font(paragraph, font_name="Times New Roman", font_size_pt=10):
    """
    Выставляет шрифт/размер именно для НУМЕРАЦИИ абзаца (цифры списков 1., 2.1, ...),
    т.к. Word берёт их из w:pPr/w:rPr.
    """
<<<<<<< HEAD
=======
    # noinspection PyProtectedMember
>>>>>>> 03ccfeb (Fix docx generation, disable PDF, stabilize template formatting)
    p = paragraph._element
    pPr = p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p.insert(0, pPr)
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)
    # шрифты
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(a), font_name)
        # убрать theme-override, если есть
        theme_attr = a + 'Theme'
        if rFonts.get(qn(theme_attr)) is not None:
            del rFonts.attrib[qn(theme_attr)]
    # размер в half-points
<<<<<<< HEAD
=======
    # noinspection DuplicatedCode
>>>>>>> 03ccfeb (Fix docx generation, disable PDF, stabilize template formatting)
    hp = str(int(font_size_pt * 2))
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), hp)
    szcs = rPr.find(qn('w:szCs'))
    if szcs is None:
        szcs = OxmlElement('w:szCs')
        rPr.append(szcs)
    szcs.set(qn('w:val'), hp)
    # baseline
    va = rPr.find(qn('w:vertAlign'))
    if va is not None:
        rPr.remove(va)
    va = OxmlElement('w:vertAlign')
    va.set(qn('w:val'), 'baseline')
    rPr.append(va)

def _normalize_paragraph(paragraph, pt: int = 10, family: str = "Times New Roman"):
    """
    Полная нормализация абзаца:
    - выставляем базовый стиль 'Normal' (Обычный), убирая спецстили
    - применяем Times New Roman 10 pt ко всем run'ам
    - отключаем 'keep_with_next' / 'keep_together', чтобы не провоцировать переносы
    """
    try:
        paragraph.style = paragraph.part.styles['Normal']
<<<<<<< HEAD
=======
    # noinspection PyBroadException
>>>>>>> 03ccfeb (Fix docx generation, disable PDF, stabilize template formatting)
    except Exception:
        # если стиль недоступен, просто продолжаем — применим прямое форматирование к run
        pass
    # формат абзаца: не склеивать с соседними
    try:
        pf = paragraph.paragraph_format
        pf.keep_with_next = False
        pf.keep_together = False
        pf.page_break_before = False
        pf.space_before = None
        pf.space_after = None
<<<<<<< HEAD
=======
    # noinspection PyBroadException
>>>>>>> 03ccfeb (Fix docx generation, disable PDF, stabilize template formatting)
    except Exception:
        pass
    _clear_paragraph_char_props(paragraph)
    _desuperscript_paragraph(paragraph)
    _ensure_para_numbering_font(paragraph, family, pt)
    # прямое форматирование на все runs
    for run in paragraph.runs:
        _force_font(run, pt, family)

_SUPERSCRIPT_MAP = str.maketrans({
    "⁰": "0", "¹": "1", "²": "2", "³": "3", "⁴": "4",
    "⁵": "5", "⁶": "6", "⁷": "7", "⁸": "8", "⁹": "9",
    "ˣ": "x", "ᵃ": "a", "ᵇ": "b", "ᶜ": "c", "ᵈ": "d", "ᵉ": "e"
})

<<<<<<< HEAD
=======
 # noinspection DuplicatedCode
>>>>>>> 03ccfeb (Fix docx generation, disable PDF, stabilize template formatting)
def _desuperscript_paragraph(paragraph):
    """
    Убирает надстрочное форматирование (x² → x2) в абзаце:
    - выключает superscript для всех runs,
    - удаляет w:vertAlign и похожие узлы,
    - переводит надстрочные цифры/символы в обычные по таблице.
    """
    # noinspection PyProtectedMember
    p = paragraph._element
    pPr = p.find(qn('w:pPr'))
    if pPr is not None:
        rPr = pPr.find(qn('w:rPr'))
        if rPr is not None:
            va = rPr.find(qn('w:vertAlign'))
            if va is not None:
                rPr.remove(va)
            # set baseline explicitly to override style inheritance
            va_new = OxmlElement('w:vertAlign')
            va_new.set(qn('w:val'), 'baseline')
            rPr.append(va_new)
    for run in paragraph.runs:
        try:
            run.font.superscript = False
            run.font.subscript = False
<<<<<<< HEAD
=======
        # noinspection PyBroadException
>>>>>>> 03ccfeb (Fix docx generation, disable PDF, stabilize template formatting)
        except Exception:
            pass
        # чистим XML выравнивание по вертикали
        # noinspection PyProtectedMember
        r = run._element
        rPr = r.find(qn('w:rPr'))
        if rPr is not None:
            va = rPr.find(qn('w:vertAlign'))
            if va is not None:
                rPr.remove(va)
        # заменяем надстрочные символы на обычные
        if run.text:
            run.text = run.text.translate(_SUPERSCRIPT_MAP)

<<<<<<< HEAD
def _replace_in_document(doc: Document, mapping: dict):
=======
def _replace_in_document(doc: DocxDocument, mapping: dict):
>>>>>>> 03ccfeb (Fix docx generation, disable PDF, stabilize template formatting)
    # Текст документа
    for para in doc.paragraphs:
        _replace_in_paragraph(para, mapping)

    # Таблицы (включая вложенные)
    def process_table(table):
        for row in table.rows:
            for cell in row.cells:
                for cell_para in cell.paragraphs:
                    _replace_in_paragraph(cell_para, mapping)
                for inner_tbl in cell.tables:
                    process_table(inner_tbl)

    for tbl in doc.tables:
        process_table(tbl)

    # Хедеры/футеры документа, если там есть плейсхолдеры
    for section in doc.sections:
        for header_para in section.header.paragraphs:
            _replace_in_paragraph(header_para, mapping)
        for header_tbl in section.header.tables:
            process_table(header_tbl)
        for footer_para in section.footer.paragraphs:
            _replace_in_paragraph(footer_para, mapping)
        for footer_tbl in section.footer.tables:
            process_table(footer_tbl)


def fill_placeholders(doc_path: Path, output_path: Path, replacements: dict):
    doc = Document(str(doc_path))
<<<<<<< HEAD
    # Enforce Normal style to TNR 10pt globally (baseline), helps avoid theme shrink
    try:
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(10)
    except Exception:
        pass
=======
    if not PRESERVE_TEMPLATE_FORMAT:
        # Enforce Normal style to TNR 10pt globally (baseline), helps avoid theme shrink
        try:
            normal_style = doc.styles['Normal']
            normal_style.font.name = 'Times New Roman'
            normal_style.font.size = Pt(10)
        except Exception:
            pass
>>>>>>> 03ccfeb (Fix docx generation, disable PDF, stabilize template formatting)
    _replace_in_document(doc, replacements)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


<<<<<<< HEAD
=======
def convert_docx_to_pdf(docx_path: Path) -> Path:
    """
    Конвертация DOCX → PDF через LibreOffice (headless).
    Требует установленный libreoffice (soffice).
    """
    out_dir = docx_path.parent
    cmd = [
        "soffice",
        "--headless",
        "--convert-to", "pdf",
        "--outdir", str(out_dir),
        str(docx_path),
    ]

    subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    return docx_path.with_suffix(".pdf")


>>>>>>> 03ccfeb (Fix docx generation, disable PDF, stabilize template formatting)
# noinspection SpellCheckingInspection,PyPep8Naming
def generate_contract_and_schedule(data: dict, out_dir: Path):
    """
    Формирует два готовых docx:
    1) Договор (templates/murabaha_template.docx)
    2) График (templates/murabaha_schedule.docx)
    Имена файлов используют data['contract_number'].
    """
    out_dir.mkdir(parents=True, exist_ok=True)

<<<<<<< HEAD
    contract_template = Path("templates/murabaha_template.docx")
    schedule_template = Path("templates/murabaha_schedule.docx")
=======
    contract_template = TEMPLATE_CONTRACT
    schedule_template = TEMPLATE_SCHEDULE
>>>>>>> 03ccfeb (Fix docx generation, disable PDF, stabilize template formatting)

    contract_out = out_dir / f"dogovor_{data['contract_number']}.docx"
    schedule_out = out_dir / f"schedule_{data['contract_number']}.docx"

    fill_placeholders(contract_template, contract_out, data)
    fill_placeholders(schedule_template, schedule_out, data)

    return contract_out, schedule_out